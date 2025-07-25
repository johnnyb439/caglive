#!/usr/bin/env python3
"""
Convert Cleared Advisory Website Audit from Markdown to Word and Excel documents.
This script reads the markdown audit file and creates:
1. A Word document with formatted content and embedded screenshots
2. An Excel document with structured audit findings
"""

import os
import re
from pathlib import Path
from datetime import datetime

# Import required packages
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import markdown
except ImportError as e:
    print(f"Missing required package: {e}")
    print("\nPlease install required packages using:")
    print("pip install python-docx openpyxl markdown")
    exit(1)

class AuditConverter:
    def __init__(self, markdown_path, screenshots_dir, output_dir):
        self.markdown_path = Path(markdown_path)
        self.screenshots_dir = Path(screenshots_dir)
        self.output_dir = Path(output_dir)
        self.content = ""
        self.sections = {}
        
    def read_markdown(self):
        """Read the markdown file content."""
        print(f"Reading markdown file: {self.markdown_path}")
        with open(self.markdown_path, 'r', encoding='utf-8') as f:
            self.content = f.read()
        print(f"✓ Successfully read {len(self.content)} characters")
        
    def parse_sections(self):
        """Parse markdown content into sections."""
        print("Parsing markdown sections...")
        
        # Split content by main headers
        lines = self.content.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            if line.startswith('## '):
                if current_section:
                    self.sections[current_section] = '\n'.join(current_content)
                current_section = line[3:].strip()
                current_content = []
            else:
                current_content.append(line)
        
        # Don't forget the last section
        if current_section:
            self.sections[current_section] = '\n'.join(current_content)
            
        print(f"✓ Parsed {len(self.sections)} sections")
        
    def create_word_document(self):
        """Create a Word document with formatted content and embedded images."""
        print("\nCreating Word document...")
        
        # Create new document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Cleared Advisory Group Website Comprehensive Audit', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph(f'Website URL: https://cleared-advisory-group-website.vercel.app/')
        doc.add_paragraph(f'Audit Date: January 22, 2025')
        doc.add_paragraph(f'Audit Type: Full Feature Analysis with Visual Documentation')
        doc.add_paragraph()
        
        # Process each section
        for section_title, section_content in self.sections.items():
            # Add section heading
            doc.add_heading(section_title, 1)
            
            # Process content line by line
            lines = section_content.strip().split('\n')
            
            for line in lines:
                # Check for image references
                img_match = re.match(r'!\[([^\]]*)\]\(([^)]+)\)', line)
                if img_match:
                    img_alt = img_match.group(1)
                    img_path = img_match.group(2)
                    
                    # Construct full image path
                    full_img_path = self.screenshots_dir / img_path.split('/')[-1]
                    
                    if full_img_path.exists():
                        # Add image caption
                        caption = doc.add_paragraph(f'Figure: {img_alt}')
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add image
                        try:
                            doc.add_picture(str(full_img_path), width=Inches(6))
                            last_paragraph = doc.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            doc.add_paragraph()  # Add spacing
                            print(f"  ✓ Embedded image: {img_path.split('/')[-1]}")
                        except Exception as e:
                            print(f"  ✗ Failed to embed image {img_path}: {e}")
                    continue
                
                # Handle different line types
                if line.startswith('### '):
                    doc.add_heading(line[4:], 2)
                elif line.startswith('**') and line.endswith('**'):
                    # Bold text
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:-2])
                    run.bold = True
                elif line.startswith('- '):
                    # Bullet point
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
                    # Numbered list
                    doc.add_paragraph(line[3:], style='List Number')
                elif line.strip():
                    # Regular paragraph
                    doc.add_paragraph(line)
        
        # Save document
        output_path = self.output_dir / 'Cleared_Advisory_Website_Audit.docx'
        doc.save(str(output_path))
        print(f"✓ Word document saved to: {output_path}")
        return output_path
        
    def create_excel_document(self):
        """Create an Excel document with structured audit findings."""
        print("\nCreating Excel document...")
        
        # Create new workbook
        wb = openpyxl.Workbook()
        
        # Define styles
        header_font = Font(bold=True, size=12, color="FFFFFF")
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        
        subheader_font = Font(bold=True, size=11)
        subheader_fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
        
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # 1. Executive Summary Sheet
        ws_summary = wb.active
        ws_summary.title = "Executive Summary"
        
        # Headers
        headers = ["Category", "Finding", "Priority", "Impact"]
        for col, header in enumerate(headers, 1):
            cell = ws_summary.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        # Summary data
        summary_data = [
            ["Overall Score", "7.2/10", "N/A", "N/A"],
            ["Key Strength", "Clear target audience definition", "N/A", "N/A"],
            ["Key Strength", "Strong value proposition for cleared professionals", "N/A", "N/A"],
            ["Key Strength", "Unique AI mock interview feature", "N/A", "N/A"],
            ["Key Strength", "Mobile-responsive design", "N/A", "N/A"],
            ["Critical Issue", "Limited interactive elements", "HIGH", "25% conversion impact"],
            ["Critical Issue", "No clear pricing information", "HIGH", "30% conversion impact"],
            ["Critical Issue", "Missing trust badges/certifications", "MEDIUM", "22% trust impact"],
            ["Critical Issue", "Limited social proof quantity", "MEDIUM", "18% engagement impact"],
        ]
        
        for row_idx, row_data in enumerate(summary_data, 2):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws_summary.cell(row=row_idx, column=col_idx, value=value)
                cell.border = border
        
        # Adjust column widths
        for col in range(1, 5):
            ws_summary.column_dimensions[get_column_letter(col)].width = 30
        
        # 2. Feature Analysis Sheet
        ws_features = wb.create_sheet("Feature Analysis")
        
        # Headers
        headers = ["Section", "Current State", "Issue", "Recommendation", "Priority", "Estimated Impact"]
        for col, header in enumerate(headers, 1):
            cell = ws_features.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        # Feature data
        feature_data = [
            ["Header & Navigation", "Clean, professional navigation", "Logo size too small", "Increase logo size by 20%", "HIGH", "15% navigation efficiency"],
            ["Header & Navigation", "8 navigation items", "No sticky navigation", "Implement sticky header", "HIGH", "15% navigation efficiency"],
            ["Hero Section", "Clear headline and CTAs", "Visual hierarchy weak", "Larger hero text", "HIGH", "25% conversion increase"],
            ["Hero Section", "Multiple CTA options", "No social proof", "Add '500+ Veterans Placed'", "HIGH", "25% conversion increase"],
            ["AI Mock Interview", "Two-tier structure", "No demo video", "Add 30-second demo", "HIGH", "40% feature adoption"],
            ["AI Mock Interview", "Unique feature", "No pricing shown", "Add pricing or 'Free Trial'", "HIGH", "40% feature adoption"],
            ["Testimonials", "Real success stories", "Only 3 testimonials", "Increase to 6-8 minimum", "MEDIUM", "22% trust increase"],
            ["Mobile Experience", "Responsive design", "Small CTAs", "Increase button sizes", "HIGH", "30% mobile conversion"],
            ["Technical SEO", "SSL implemented", "Missing meta descriptions", "Add comprehensive meta tags", "HIGH", "35% organic traffic"],
            ["Footer", "Contact info present", "No social media links", "Add social media icons", "MEDIUM", "12% navigation improvement"],
        ]
        
        for row_idx, row_data in enumerate(feature_data, 2):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws_features.cell(row=row_idx, column=col_idx, value=value)
                cell.border = border
                cell.alignment = Alignment(wrap_text=True, vertical="top")
        
        # Adjust column widths
        ws_features.column_dimensions['A'].width = 20
        ws_features.column_dimensions['B'].width = 30
        ws_features.column_dimensions['C'].width = 30
        ws_features.column_dimensions['D'].width = 35
        ws_features.column_dimensions['E'].width = 12
        ws_features.column_dimensions['F'].width = 25
        
        # 3. Recommendations Sheet
        ws_recs = wb.create_sheet("Recommendations")
        
        # Headers
        headers = ["Timeframe", "Category", "Recommendation", "Expected Outcome", "Resources Needed"]
        for col, header in enumerate(headers, 1):
            cell = ws_recs.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        # Recommendations data
        recs_data = [
            ["Week 1", "Quick Win", "Add 'Free Consultation' badge", "10% CTR increase", "Designer - 1 hour"],
            ["Week 1", "Quick Win", "Implement exit-intent popup", "15% conversion increase", "Developer - 3 hours"],
            ["Week 1", "Quick Win", "Add trust badges", "20% trust increase", "Designer - 2 hours"],
            ["Week 2", "Trust Building", "Add more testimonials", "22% trust increase", "Content team - 4 hours"],
            ["Week 2", "Trust Building", "Implement live chat", "25% engagement increase", "Developer - 8 hours"],
            ["Month 2", "Content", "Create resource center", "30% engagement increase", "Content team - 40 hours"],
            ["Month 2", "Content", "Audience landing pages", "35% conversion increase", "Full team - 60 hours"],
            ["Month 3", "Advanced", "Launch A/B testing", "Continuous improvement", "Analytics team - ongoing"],
            ["Month 3", "Advanced", "Mobile app MVP", "40% user retention", "Dev team - 200 hours"],
        ]
        
        for row_idx, row_data in enumerate(recs_data, 2):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws_recs.cell(row=row_idx, column=col_idx, value=value)
                cell.border = border
                cell.alignment = Alignment(wrap_text=True, vertical="top")
        
        # Adjust column widths
        for col in range(1, 6):
            ws_recs.column_dimensions[get_column_letter(col)].width = 25
        
        # 4. Implementation Roadmap Sheet
        ws_roadmap = wb.create_sheet("Implementation Roadmap")
        
        # Headers
        headers = ["Phase", "Timeline", "Tasks", "Success Metrics", "Budget Estimate"]
        for col, header in enumerate(headers, 1):
            cell = ws_roadmap.cell(row=1, column=col, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = border
        
        # Roadmap data
        roadmap_data = [
            ["Phase 1: Critical Fixes", "Week 1-2", "Hero section, CTAs, Navigation", "25% conversion increase", "$2,000-$3,000"],
            ["Phase 2: Trust Building", "Week 3-4", "Testimonials, Badges, Live Chat", "22% trust increase", "$3,000-$4,000"],
            ["Phase 3: Content Enhancement", "Month 2", "Resource Center, Landing Pages", "35% engagement increase", "$8,000-$10,000"],
            ["Phase 4: Advanced Features", "Month 3", "A/B Testing, Mobile App", "40% overall improvement", "$15,000-$20,000"],
        ]
        
        for row_idx, row_data in enumerate(roadmap_data, 2):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws_roadmap.cell(row=row_idx, column=col_idx, value=value)
                cell.border = border
                cell.alignment = Alignment(wrap_text=True, vertical="top")
                
                # Highlight phases with different colors
                if col_idx == 1:
                    if "Phase 1" in value:
                        cell.fill = PatternFill(start_color="FFE4E1", end_color="FFE4E1", fill_type="solid")
                    elif "Phase 2" in value:
                        cell.fill = PatternFill(start_color="E1F5FE", end_color="E1F5FE", fill_type="solid")
                    elif "Phase 3" in value:
                        cell.fill = PatternFill(start_color="F3E5F5", end_color="F3E5F5", fill_type="solid")
                    elif "Phase 4" in value:
                        cell.fill = PatternFill(start_color="E8F5E9", end_color="E8F5E9", fill_type="solid")
        
        # Adjust column widths
        for col in range(1, 6):
            ws_roadmap.column_dimensions[get_column_letter(col)].width = 25
        
        # Save workbook
        output_path = self.output_dir / 'Cleared_Advisory_Website_Audit.xlsx'
        wb.save(str(output_path))
        print(f"✓ Excel document saved to: {output_path}")
        return output_path

def main():
    """Main function to run the conversion."""
    print("=== Cleared Advisory Website Audit Converter ===\n")
    
    # Define paths
    markdown_path = "/Users/tone/cleared-advisory/Cleared_Advisory_Website_Audit_Complete.md"
    screenshots_dir = "/Users/tone/cleared-advisory/website_audit_screenshots"
    output_dir = Path.home() / "Desktop"
    
    # Create converter instance
    converter = AuditConverter(markdown_path, screenshots_dir, output_dir)
    
    try:
        # Read and parse markdown
        converter.read_markdown()
        converter.parse_sections()
        
        # Create documents
        word_path = converter.create_word_document()
        excel_path = converter.create_excel_document()
        
        print("\n=== Conversion Complete! ===")
        print(f"\nDocuments saved to desktop:")
        print(f"  • Word: {word_path}")
        print(f"  • Excel: {excel_path}")
        
        # Display file sizes
        word_size = word_path.stat().st_size / 1024 / 1024  # MB
        excel_size = excel_path.stat().st_size / 1024  # KB
        print(f"\nFile sizes:")
        print(f"  • Word: {word_size:.2f} MB")
        print(f"  • Excel: {excel_size:.1f} KB")
        
    except Exception as e:
        print(f"\n✗ Error during conversion: {e}")
        import traceback
        traceback.print_exc()
        return 1
    
    return 0

if __name__ == "__main__":
    exit(main())