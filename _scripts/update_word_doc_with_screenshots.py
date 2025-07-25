#!/usr/bin/env python3
import os
from docx import Document
from docx.shared import Inches
from PIL import Image
import io

def update_word_document_with_screenshots():
    """Update the Word document with the new properly captured screenshots"""
    
    # Load the existing Word document
    doc_path = "/Users/tone/Desktop/Cleared_Advisory_Website_Audit.docx"
    doc = Document(doc_path)
    
    # Create a new document with updated screenshots
    new_doc = Document()
    
    # Copy styles from original document
    for style in doc.styles:
        try:
            if style.name not in new_doc.styles:
                new_doc.styles.add_style(style.name, style.type)
        except:
            pass
    
    # Screenshot mapping - which screenshot to use for each section
    screenshot_mapping = {
        "Header Desktop": "/Users/tone/Desktop/website_audit_screenshots_v2/02_header_desktop.png",
        "Homepage Full Desktop": "/Users/tone/Desktop/website_audit_screenshots_v2/01_homepage_full_desktop.png",
        "Target Audience": "/Users/tone/Desktop/website_audit_screenshots_v2/04_section_2_desktop.png",
        "Mock Interview Section": "/Users/tone/Desktop/website_audit_screenshots_v2/04_section_4_desktop.png",
        "Testimonials Section": "/Users/tone/Desktop/website_audit_screenshots_v2/04_section_5_desktop.png",
        "Footer Section": "/Users/tone/Desktop/website_audit_screenshots_v2/02_footer_desktop.png",
        "Mobile Homepage": "/Users/tone/Desktop/website_audit_screenshots_v2/01_homepage_full_mobile.png",
        "Mobile Menu": "/Users/tone/Desktop/website_audit_screenshots_v2/03_mobile_menu_expanded.png"
    }
    
    # Process each element in the original document
    for element in doc.element.body:
        if element.tag.endswith('p'):
            # Copy paragraph
            paragraph = doc.paragraphs[doc.element.body.index(element)]
            new_p = new_doc.add_paragraph()
            new_p.text = paragraph.text
            
            # Copy paragraph formatting
            if paragraph.style:
                new_p.style = paragraph.style
            new_p.alignment = paragraph.alignment
            
            # Copy run formatting
            if paragraph.runs:
                new_p.runs[0].bold = paragraph.runs[0].bold if paragraph.runs else False
                new_p.runs[0].italic = paragraph.runs[0].italic if paragraph.runs else False
                if paragraph.runs and paragraph.runs[0].font.size:
                    new_p.runs[0].font.size = paragraph.runs[0].font.size
                    
        elif element.tag.endswith('tbl'):
            # Copy table
            table_index = [i for i, e in enumerate(doc.element.body) if e.tag.endswith('tbl')].index(doc.element.body.index(element))
            if table_index < len(doc.tables):
                old_table = doc.tables[table_index]
                new_table = new_doc.add_table(rows=len(old_table.rows), cols=len(old_table.columns))
                
                # Copy table content
                for i, row in enumerate(old_table.rows):
                    for j, cell in enumerate(row.cells):
                        new_table.cell(i, j).text = cell.text
    
    # Now add screenshots at appropriate positions
    print("Adding screenshots to the document...")
    
    # Add screenshots based on section headings
    for i, paragraph in enumerate(new_doc.paragraphs):
        # Header section
        if "Header & Navigation Analysis" in paragraph.text and paragraph.style.name.startswith('Heading'):
            # Add header screenshot after this heading
            new_doc.paragraphs[i].add_run().add_break()
            try:
                new_doc.add_picture(screenshot_mapping["Header Desktop"], width=Inches(6))
                new_doc.add_paragraph("Figure: Header and Navigation", style='Caption')
                print("Added header screenshot")
            except Exception as e:
                print(f"Error adding header screenshot: {e}")
                
        # Hero section 
        elif "Hero Section" in paragraph.text and paragraph.style.name.startswith('Heading'):
            new_doc.paragraphs[i].add_run().add_break()
            try:
                # Crop the hero section from the full page screenshot
                img = Image.open(screenshot_mapping["Homepage Full Desktop"])
                # Crop to show just the hero section (top portion)
                width, height = img.size
                cropped = img.crop((0, 0, width, int(height * 0.3)))
                
                # Save to bytes
                img_bytes = io.BytesIO()
                cropped.save(img_bytes, format='PNG')
                img_bytes.seek(0)
                
                new_doc.add_picture(img_bytes, width=Inches(6))
                new_doc.add_paragraph("Figure: Hero Section with CTAs", style='Caption')
                print("Added hero section screenshot")
            except Exception as e:
                print(f"Error adding hero screenshot: {e}")
                
        # Target Audience
        elif "Target Audience Section" in paragraph.text and paragraph.style.name.startswith('Heading'):
            new_doc.paragraphs[i].add_run().add_break()
            try:
                new_doc.add_picture(screenshot_mapping["Target Audience"], width=Inches(6))
                new_doc.add_paragraph("Figure: Target Audience Section", style='Caption')
                print("Added target audience screenshot")
            except Exception as e:
                print(f"Error adding target audience screenshot: {e}")
                
        # Mock Interview
        elif "AI Mock Interview Feature" in paragraph.text and paragraph.style.name.startswith('Heading'):
            new_doc.paragraphs[i].add_run().add_break()
            try:
                new_doc.add_picture(screenshot_mapping["Mock Interview Section"], width=Inches(6))
                new_doc.add_paragraph("Figure: AI Mock Interview Feature", style='Caption')
                print("Added mock interview screenshot")
            except Exception as e:
                print(f"Error adding mock interview screenshot: {e}")
                
        # Testimonials
        elif "Success Stories/Testimonials" in paragraph.text and paragraph.style.name.startswith('Heading'):
            new_doc.paragraphs[i].add_run().add_break()
            try:
                new_doc.add_picture(screenshot_mapping["Testimonials Section"], width=Inches(6))
                new_doc.add_paragraph("Figure: Success Stories and Testimonials", style='Caption')
                print("Added testimonials screenshot")
            except Exception as e:
                print(f"Error adding testimonials screenshot: {e}")
                
        # Footer
        elif "Footer Analysis" in paragraph.text and paragraph.style.name.startswith('Heading'):
            new_doc.paragraphs[i].add_run().add_break()
            try:
                new_doc.add_picture(screenshot_mapping["Footer Section"], width=Inches(6))
                new_doc.add_paragraph("Figure: Footer Section", style='Caption')
                print("Added footer screenshot")
            except Exception as e:
                print(f"Error adding footer screenshot: {e}")
                
        # Mobile Responsiveness
        elif "Mobile Responsiveness" in paragraph.text and paragraph.style.name.startswith('Heading'):
            new_doc.paragraphs[i].add_run().add_break()
            try:
                # Add mobile screenshots side by side
                new_doc.add_picture(screenshot_mapping["Mobile Homepage"], width=Inches(2.5))
                new_doc.add_paragraph("Figure: Mobile Homepage View", style='Caption')
                
                new_doc.add_picture(screenshot_mapping["Mobile Menu"], width=Inches(2.5))
                new_doc.add_paragraph("Figure: Mobile Menu Expanded", style='Caption')
                print("Added mobile screenshots")
            except Exception as e:
                print(f"Error adding mobile screenshots: {e}")
    
    # Save the updated document
    output_path = "/Users/tone/Desktop/Cleared_Advisory_Website_Audit_Updated.docx"
    new_doc.save(output_path)
    print(f"\nUpdated Word document saved to: {output_path}")
    
    # Get file size
    file_size = os.path.getsize(output_path)
    print(f"File size: {file_size / (1024*1024):.2f} MB")

if __name__ == "__main__":
    # Install required packages if needed
    import subprocess
    subprocess.run(["python3", "-m", "pip", "install", "python-docx", "Pillow"], capture_output=True)
    
    update_word_document_with_screenshots()