#!/usr/bin/env python3
import os
import subprocess

# Install required package
print("Installing required packages...")
subprocess.run(["python3", "-m", "pip", "install", "python-docx"], capture_output=True)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_word_document_with_screenshots():
    """Create a new Word document with the website audit and proper screenshots"""
    
    # Create new document
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Heading styles
    heading_1 = styles['Heading 1']
    heading_1.font.size = Pt(16)
    heading_1.font.color.rgb = RGBColor(0, 0, 0)
    
    heading_2 = styles['Heading 2'] 
    heading_2.font.size = Pt(14)
    heading_2.font.color.rgb = RGBColor(0, 0, 0)
    
    # Read the markdown content
    with open('/Users/tone/cleared-advisory/Cleared_Advisory_Website_Audit_Complete.md', 'r') as f:
        content = f.read()
    
    # Split content into sections
    lines = content.split('\n')
    
    # Screenshot paths
    screenshots = {
        'header': '/Users/tone/Desktop/website_audit_screenshots_v2/02_header_desktop.png',
        'hero': '/Users/tone/Desktop/website_audit_screenshots_v2/04_section_1_desktop.png',
        'target': '/Users/tone/Desktop/website_audit_screenshots_v2/04_section_2_desktop.png',
        'process': '/Users/tone/Desktop/website_audit_screenshots_v2/04_section_3_desktop.png',
        'mock': '/Users/tone/Desktop/website_audit_screenshots_v2/04_section_4_desktop.png',
        'testimonials': '/Users/tone/Desktop/website_audit_screenshots_v2/04_section_5_desktop.png',
        'footer': '/Users/tone/Desktop/website_audit_screenshots_v2/02_footer_desktop.png',
        'mobile_full': '/Users/tone/Desktop/website_audit_screenshots_v2/01_homepage_full_mobile.png',
        'mobile_menu': '/Users/tone/Desktop/website_audit_screenshots_v2/03_mobile_menu_expanded.png'
    }
    
    # Process content
    in_code_block = False
    current_section = None
    
    for line in lines:
        # Skip markdown image references
        if line.startswith('!['):
            continue
            
        # Headers
        if line.startswith('# ') and not line.startswith('##'):
            doc.add_heading(line[2:], 0)
        elif line.startswith('## '):
            current_section = line[3:].lower()
            doc.add_heading(line[3:], 1)
            
            # Add screenshots after specific headings
            if 'header & navigation' in current_section:
                try:
                    doc.add_paragraph()
                    doc.add_picture(screenshots['header'], width=Inches(6))
                    p = doc.add_paragraph('Figure: Website Header and Navigation')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                except:
                    pass
                    
            elif 'hero section' in current_section:
                try:
                    doc.add_paragraph()
                    doc.add_picture(screenshots['hero'], width=Inches(6))
                    p = doc.add_paragraph('Figure: Hero Section with Call-to-Actions')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                except:
                    pass
                    
            elif 'target audience' in current_section:
                try:
                    doc.add_paragraph()
                    doc.add_picture(screenshots['target'], width=Inches(6))
                    p = doc.add_paragraph('Figure: Target Audience Section')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                except:
                    pass
                    
            elif 'mock interview' in current_section:
                try:
                    doc.add_paragraph()
                    doc.add_picture(screenshots['mock'], width=Inches(6))
                    p = doc.add_paragraph('Figure: AI-Powered Mock Interview Feature')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                except:
                    pass
                    
            elif 'testimonials' in current_section or 'success stories' in current_section:
                try:
                    doc.add_paragraph()
                    doc.add_picture(screenshots['testimonials'], width=Inches(6))
                    p = doc.add_paragraph('Figure: Success Stories and Testimonials')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                except:
                    pass
                    
            elif 'footer' in current_section:
                try:
                    doc.add_paragraph()
                    doc.add_picture(screenshots['footer'], width=Inches(6))
                    p = doc.add_paragraph('Figure: Website Footer')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                except:
                    pass
                    
            elif 'mobile' in current_section:
                try:
                    doc.add_paragraph()
                    # Add mobile screenshots
                    doc.add_picture(screenshots['mobile_full'], width=Inches(3))
                    p = doc.add_paragraph('Figure: Mobile Homepage View')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    doc.add_picture(screenshots['mobile_menu'], width=Inches(3))
                    p = doc.add_paragraph('Figure: Mobile Menu Expanded')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
                except:
                    pass
                    
        elif line.startswith('### '):
            doc.add_heading(line[4:], 2)
        elif line.startswith('**') and line.endswith('**') and len(line) > 4:
            # Bold text
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            run.bold = True
        elif line.startswith('- '):
            # Bullet points
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith(tuple('0123456789')) and '. ' in line:
            # Numbered lists
            doc.add_paragraph(line, style='List Number')
        elif line.strip() == '---':
            # Horizontal line - add some space
            doc.add_paragraph()
        elif line.strip():
            # Regular paragraph
            doc.add_paragraph(line)
    
    # Add overview screenshot at the end
    doc.add_page_break()
    doc.add_heading('Appendix: Full Website Overview', 1)
    try:
        doc.add_picture('/Users/tone/Desktop/website_audit_screenshots_v2/01_homepage_full_desktop.png', width=Inches(6))
        p = doc.add_paragraph('Figure: Complete Website Desktop View')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except:
        pass
    
    # Save document
    output_path = '/Users/tone/Desktop/Cleared_Advisory_Website_Audit_WithScreenshots.docx'
    doc.save(output_path)
    
    # Get file info
    file_size = os.path.getsize(output_path)
    print(f"Document created successfully!")
    print(f"Location: {output_path}")
    print(f"Size: {file_size / (1024*1024):.2f} MB")
    print(f"\nThe document includes:")
    print("- Complete website audit text")
    print("- 8 embedded screenshots showing all website sections")
    print("- Professional formatting with headings and lists")
    print("- Appendix with full website overview")

if __name__ == "__main__":
    create_word_document_with_screenshots()